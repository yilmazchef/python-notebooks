import os
import sys
import time
import shutil as sh
import watchdog.events
import watchdog.observers

import configparser
import json

from uuid import uuid4
from urllib.parse import quote

import collections.abc
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from progressbar import AnimatedMarker, Bar, BouncingBar, Counter, ETA, \
    AdaptiveETA, FileTransferSpeed, FormatLabel, Percentage, \
    ProgressBar, ReverseBar, RotatingMarker, \
    SimpleProgress, Timer, UnknownLength
    

def scrape_website_in_markdown(url: str, output_file: str) -> str:
    """Scrapes a website and saves it as a markdown file

    Args:
        url (str): the url of the website
        output_file (str): the path to the markdown file

    Returns:
        str: the path to the markdown file
    """
    cmdlet = f"wget -O \"{output_file}\" \"{url}\""
    os.system(cmdlet)
    return output_file

def get_files(src_path: str, file_ext: str) -> list:

    source_file_list = []

    for root, dirs, files in os.walk(src_path):
        for file in files:
            if file.endswith(file_ext):
                source_file_list.append(os.path.join(root, file))

    return source_file_list


def move_file(file_path: str, dest_folder: str):

    if not os.path.exists(dest_folder):
        os.mkdir(dest_folder)

    if(os.path.exists(os.path.join(dest_folder, os.path.basename(file_path)))):
        os.remove(os.path.join(dest_folder, os.path.basename(file_path)))
    
    sh.move(file_path, dest_folder)
    
    return os.path.join(dest_folder, os.path.basename(file_path))

def html_to_markdown(html_file: str) -> str:
    """Converts an html file to a markdown file

    Args:
        html_file (str): the path to the html file

    Returns:
        str: the path to the markdown file
    """
    output_file = html_file.replace(".html", ".md")
    cmdlet = f"pandoc -f html -t markdown \"{html_file}\" -o \"{output_file}\""
    
    os.system(cmdlet)
    
    # get the parent folder name
    parent_folder = os.path.basename(os.path.dirname(html_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "MDs", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    md_file = move_file(md_file, target_folder)
    
    return output_file


def markdown_to_html(md_file: str) -> str:

    html_file = md_file.replace(".md", ".html")
    cmdlet = f"pandoc -o \"{html_file}\" \"{md_file}\""
    os.system(cmdlet)
    
    # get the parent folder name
    parent_folder = os.path.basename(os.path.dirname(md_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "HTMLs", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
        
    html_file = move_file(html_file, target_folder)
    
    return html_file


def image_to_base64(image_file):

    # need base 64
    import base64
    import sys
    # open the image
    image = open(image_file, 'rb')
    # read it
    image_read = image.read()
    # encode it as base 64
    # after python>=3.9, use `encodebytes` instead of `encodestring`
    image_64_encode = base64.encodestring(image_read) if sys.version_info < (
        3, 9) else base64.encodebytes(image_read)
    # convert the image base 64 to a string
    image_string = str(image_64_encode)
    # replace the newline characters
    image_string = image_string.replace("\\n", "")
    # replace the initial binary
    image_string = image_string.replace("b'", "")
    # replace the final question mark
    image_string = image_string.replace("'", "")
    # add the image tags
    image_string = '<p><img src="data:image/png;base64,' + image_string + '"></p>'
    # write it out
    output_file = image_file.replace(".png", ".html")
    
    image_result = open(output_file, 'w')
    image_result.write(image_string)


def notebook_to_markdown(ipynb_file: str) -> str:
    """
    Converts an ipynb file to a markdown file
    @param ipynb_file: the path to the ipynb file
    @return: the path to the markdown file
    """
    
    md_temp_file = ipynb_file.replace(".ipynb", ".md")
    # create the command to convert the ipynb file to markdown
    cmdlet = f"jupyter nbconvert --to markdown \"{ipynb_file}\""
    # run the command using the os module function system() but do not print the output
    os.system(cmdlet)
    
    # get the parent folder name
    parent_folder = os.path.basename(os.path.dirname(ipynb_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "MDs", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    if(os.path.exists(md_temp_file)):
        md_file = move_file(md_temp_file, target_folder)
    
    # return the path to the markdown file
    return md_file

def notebook_to_python_script(ipynb_file: str) -> str:
    """
    Converts an ipynb file to a python script
    @param ipynb_file: the path to the ipynb file
    @return: the path to the python script
    """
    
    py_temp_file = ipynb_file.replace(".ipynb", ".py")
    
    # create the command to convert the ipynb file to python script
    cmdlet = f"jupyter nbconvert --to script \"{ipynb_file}\""
    # run the command using the os module function system() but do not print the output
    os.system(cmdlet)
    
    # get the parent folder name    
    parent_folder = os.path.basename(os.path.dirname(ipynb_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "Scripts", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    if(os.path.exists(py_temp_file)):
        py_file = move_file(py_temp_file, target_folder)
    
    # return the path to the python script
    return py_file


def markdown_to_notebook(md_file: str) -> str:
    """_summary_

    Args:
        md_file (str): _description_

    Returns:
        str: _description_
    """
    ipynb_file = md_file.replace(".md", ".ipynb")
    cmdlet = f"pandoc \"{md_file}\" -o \"{ipynb_file}\""
    
    os.system(cmdlet)
    
    # get the parent folder name    
    parent_folder = os.path.basename(os.path.dirname(md_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "Notebooks", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    ipynb_file = move_file(md_file.replace(".md", ".ipynb"), target_folder)
    
    return ipynb_file


def markdown_to_open_document(md_file: str):
    """
    Converts a markdown file to an open document file
    @param md_file: the path to the markdown file
    @return: the path to the open document file
    """
    # create the command to convert the markdown file to open document
    cmdlet = f"pandoc \"{md_file}\" -o \"{md_file.replace('.md', '.odt')}\""
    # run the command using the os module function system() but do not print the output
    os.system(cmdlet)
    
    # get the parent folder name    
    parent_folder = os.path.basename(os.path.dirname(md_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "ODTs", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    odt_file = move_file(md_file.replace(".ipynb", ".odt"), target_folder)
    
    # return the path to the open document file
    return odt_file


def markdown_to_powerpoint(md_file: str) -> str:

    c = collections
    c.abc = collections.abc

    cmdlet = f"pandoc -V fontsize=12pt \"{md_file}\" -s --wrap auto -o \"{md_file.replace('.md', '.pptx')}\""
    
    os.system(cmdlet)
    
    # get the parent folder name
    parent_folder = os.path.basename(os.path.dirname(md_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "PPTs", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    # move the converted files to the respective folders
    pptx_file = move_file(md_file.replace(".md", ".pptx"), target_folder)

    return pptx_file


def markdown_to_word_document(md_file: str) -> str:

    docx_temp_file = md_file.replace(".md", ".docx")
    cmdlet = f"pandoc \"{md_file}\" -f markdown -o \"{docx_temp_file}\""
    os.system(cmdlet)
    
    add_header_to_word_document(docx_temp_file, os.path.join(os.getcwd(), "Templates", "header.png"), open(
        os.path.join(os.getcwd(), "templates", "header.txt"), 'r'))
    add_footer_to_word_document(docx_temp_file, os.path.join(os.getcwd(), "Templates", "footer.png"), open(
        os.path.join(os.getcwd(), "templates", "footer.txt"), 'r'))
    
    # get the parent folder name
    parent_folder = os.path.basename(os.path.dirname(md_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "DOCs", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
        
    if(os.path.exists(docx_temp_file)):
        docx_file = move_file(docx_file, target_folder)
    
    return docx_file


def add_image_to_word_document(fileref, question, tpl, width=None):

    if fileref.__class__.__name__ in ('DAFile', 'DAFileList', 'DAFileCollection', 'DALocalFile', 'DAStaticFile'):
        file_info = dict(fullpath=fileref.path())
    else:
        file_info = server.file_finder(fileref, question=question)
    if 'path' in file_info and 'extension' in file_info:
        docassemble.base.filter.convert_svg_to_png(file_info)
    if 'fullpath' not in file_info:
        return '[FILE NOT FOUND]'
    if width is not None:
        m = re.search(r'^([0-9\.]+) *([A-Za-z]*)', str(width))
        if m:
            amount = float(m.group(1))
            units = m.group(2).lower()
            if units in ['in', 'inches', 'inch']:
                the_width = Inches(amount)
            elif units in ['pt', 'pts', 'point', 'points']:
                the_width = Pt(amount)
            elif units in ['mm', 'millimeter', 'millimeters']:
                the_width = Mm(amount)
            elif units in ['cm', 'centimeter', 'centimeters']:
                the_width = Cm(amount)
            elif units in ['twp', 'twip', 'twips']:
                the_width = Twips(amount)
            else:
                the_width = Pt(amount)
        else:
            the_width = Inches(2)
    else:
        the_width = Inches(2)
    return InlineImage(tpl, file_info['fullpath'], the_width)


def add_header_to_word_document(docx_file, header_image, header_text=None):

    document = None

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        header = document.sections[0].header

        htable = header.add_table(1, 2, Inches(4))
        htab_cells = htable.rows[0].cells
        ht0 = htab_cells[0].add_paragraph()
        kh = ht0.add_run()
        kh.add_picture(header_image, width=Inches(3))

        if header_text is not None:
            ht1 = htab_cells[1].add_paragraph(header_text)
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # font size of header text
            ht1.style.font.size = Pt(10)

            # saving the blank document
        document.save(docx_file)


def add_footer_to_word_document(docx_file, footer_image, footer_text=None):

    document = None

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        footer = document.sections[0].footer

        ftable = footer.add_table(1, 2, Inches(4))
        ftab_cells = ftable.rows[0].cells
        ft0 = ftab_cells[0].add_paragraph()
        fh = ft0.add_run()

        imgSize = Inches(4)

        if footer_text is not None:
            imgSize = Inches(4.5)
            ft1 = ftab_cells[1].add_paragraph(footer_text)
            ft1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ft1.style.font.size = Pt(10)

        fh.add_picture(footer_image, width=imgSize)

        document.save(docx_file)


def markdown_to_pdf(md_file: str) -> str:

    pdf_file = md_file.replace(".md", ".pdf")
    cmdlet = f"pandoc \"{md_file}\" -o \"{pdf_file}\""
    os.system(cmdlet)
    
    # get the parent folder name
    parent_folder = os.path.basename(os.path.dirname(md_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "PDFs", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    # move the converted files to the respective folders
    pdf_file = move_file(md_file.replace(".md", ".pdf"), target_folder)
    
    return pdf_file


def markdown_to_ebook(md_file: str) -> str:

    epub_file = md_file.replace(".md", ".epub")
    cmdlet = f"pandoc \"{md_file}\" -o \"{epub_file}\""
    os.system(cmdlet)

    # get the parent folder name
    parent_folder = os.path.basename(os.path.dirname(md_file))
    
    # move the converted file to the corresponding folder
    target_folder = os.path.join(os.getcwd(), ".synced", "eBooks", parent_folder)
    
    # if the folder does not exist, create it
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    # move the converted files to the respective folders
    epub_file = move_file(md_file.replace(".md", ".epub"), target_folder)    

    return epub_file


def reset_progress_bar(max_value: int):
    widgets = [Percentage(),
               ' ', Bar(),
               ' ', ETA(),
               ' ', AdaptiveETA()]
    return ProgressBar(widgets=widgets, maxval=max_value)


def update_all():

    src_path = sys.argv[1] if len(sys.argv) > 1 else os.getcwd()
    notebook_file_list = get_files(src_path, ".ipynb")
    
    pBar = reset_progress_bar(len(notebook_file_list))
    pBar.start()
    pBarCount = 0

    for notebook_file in notebook_file_list:
        notebook_to_markdown(notebook_file)
        notebook_to_python_script(notebook_file)
        
        pBarCount += 1
        pBar.update(pBarCount)

    pBar.finish()

    source_file_list = get_files(os.path.join(os.getcwd(), ".synced", "MDs"), ".md")

    pBar = reset_progress_bar(len(source_file_list))
    pBar.start()
    pBarCount = 0

    for source_file in source_file_list:
        
        markdown_to_pdf(source_file)
        markdown_to_word_document(source_file)
        # markdown_to_powerpoint(source_file)

        pBarCount += 1
        pBar.update(pBarCount)

    pBar.finish()


if __name__ == "__main__":
    update_all()
