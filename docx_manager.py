import os
from importlib.metadata import PackageNotFoundError
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT


def addHeader(docxPath, headerImage, headerText=None):

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docxPath}"):

        # opening the existing document
        document = Document(docxPath)

        header = document.sections[0].header

        htable = header.add_table(1, 2, Inches(4))
        htab_cells = htable.rows[0].cells
        ht0 = htab_cells[0].add_paragraph()
        kh = ht0.add_run()
        kh.add_picture(headerImage, width=Inches(3))

        if headerText is not None:
            ht1 = htab_cells[1].add_paragraph(headerText)
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
                # saving the blank document
        document.save(docxPath)


def addFooter(docxPath, footerImage, footerText=None):

    document = None

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docxPath}"):

        # opening the existing document
        document = Document(docxPath)

        footer = document.sections[0].footer

        ftable = footer.add_table(1, 2, Inches(4))
        ftab_cells = ftable.rows[0].cells
        ft0 = ftab_cells[0].add_paragraph()
        fh = ft0.add_run()

        imgSize = Inches(4)

        if footerText is not None:
            imgSize = Inches(3)
            ft1 = ftab_cells[1].add_paragraph(footerText)
            ft1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fh.add_picture(footerImage, width=imgSize)

        document.save(docxPath)


if __name__ == "__main__":

    docxPath = "intec_template.docx"

    text = """
    Rouppeplein 16
    1000 Brussel
    Tel. 02 411 29 07
    ondernemingsnr. 0475319893
    RPR BRUSSEL-NEDERLANDSTALIG
    wouter.vandenberge@intecbrussel.be                                
    """

    image = "Temp/IntecHeader.png"

    addHeader(docxPath, image, text)

    text = "Some footer text.."

    image = "Temp/IntecFooter.png"

    addFooter(docxPath, image)
